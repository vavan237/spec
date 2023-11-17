

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK

from docxtpl import DocxTemplate
import docx
from docx2pdf import convert


def create_receipt(qr, name, adress, complete_date, period, ls, tax, part, total, summa, summ_pay):
    doc = DocxTemplate('check-final.docx')
    context = {'qr': qr,
                'fio': name,
                'adres': adress,
                'complete_date': complete_date,
                'period': period,
                'ls': ls,
                'tax': tax,
                'part': part,
                'total': total,
                'summa': summa,
                'summ_pay': summ_pay

               }
    doc.render(context)
    doc.save('check-final.docx')

def add_receipt():
    doc = Document('check-final.docx')
    doc.add_paragraph('{{key}}')
    doc.save('check-final.docx')
    tpl = DocxTemplate('check-final.docx')
    sd = tpl.new_subdoc('reciept.docx')
    context = {'key': sd}
    tpl.render(context)
    tpl.save('check-final.docx')

def create_file():
    doc = docx.Document()
    doc.save('check-final.docx')

def c():
    doc = docx.Document()
    tpl = DocxTemplate(doc)
    sd = tpl.new_subdoc('reciept.docx')
    p = doc.add_paragraph()
    #tbl=doc2.tables[0]._tbl
    #p.add_run().element.addnext(tbl)
    doc.save('check-final.docx')



#convert("check-final.docx", "new-check.pdf")
